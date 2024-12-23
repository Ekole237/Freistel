from datetime import datetime
from pathlib import Path
import os
from typing import Dict, Any, List, Tuple
import matplotlib.pyplot as plt
import numpy as np
from openai import OpenAI
from dotenv import load_dotenv
import seaborn as sns
from scipy.stats import chi2
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import markdown
import re
import hashlib
import json

# Charger les variables d'environnement
load_dotenv()

class ReportGenerator:
    def __init__(self, output_dir: str = "reports"):
        """Initialise le générateur de rapports."""
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.cache_dir = self.output_dir / "cache"
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        
    def generate_cache_key(self, data: Dict[str, Any]) -> str:
        """Génère une clé de cache unique basée sur les données d'analyse."""
        # On ne prend que les données pertinentes pour les visualisations
        cache_data = {
            'block_size': data['block_size'],
            'num_rounds': data['num_rounds'],
            'differences': data['differences'],
            'probability': data['probability'],
            'input_diff': data['input_diff'],
            'output_diff': data['output_diff']
        }
        # Convertir en JSON et calculer le hash
        data_str = json.dumps(cache_data, sort_keys=True)
        return hashlib.sha256(data_str.encode()).hexdigest()[:12]
    
    def get_cached_images(self, cache_key: str) -> Dict[str, Path]:
        """Récupère les images en cache si elles existent."""
        image_types = ['probability', 'correlation', 'boxplot', 'convergence']
        cached_images = {}
        
        for img_type in image_types:
            cached_path = self.cache_dir / f"{cache_key}_{img_type}.png"
            if cached_path.exists():
                cached_images[img_type] = cached_path
                
        return cached_images
    
    def save_to_cache(self, cache_key: str, images: Dict[str, Path]):
        """Sauvegarde les images dans le cache."""
        for img_type, img_path in images.items():
            cache_path = self.cache_dir / f"{cache_key}_{img_type}.png"
            if img_path.exists():
                img_path.rename(cache_path)

    def setup_plotting_style(self):
        """Configure un style de graphique professionnel et cohérent."""
        try:
            import seaborn as sns
            
            # Définition de la palette de couleurs personnalisée
            colors = {
                'primary': '#2C3E50',      # Bleu foncé
                'secondary': '#E74C3C',    # Rouge
                'accent': '#3498DB',       # Bleu clair
                'neutral': '#95A5A6',      # Gris
                'success': '#27AE60',      # Vert
                'warning': '#F39C12',      # Orange
                'background': '#ECF0F1',   # Gris clair
                'grid': '#BDC3C7'          # Gris moyen
            }
            
            # Création d'une palette personnalisée
            sns.set_palette([
                colors['primary'],
                colors['secondary'],
                colors['accent'],
                colors['success'],
                colors['warning']
            ])
            
            # Configuration du style de base
            sns.set_style("whitegrid", {
                'axes.grid': True,
                'grid.color': colors['grid'],
                'grid.linestyle': '--',
                'grid.alpha': 0.3,
                'axes.edgecolor': colors['primary'],
                'axes.linewidth': 1.5,
                'axes.facecolor': colors['background']
            })
            
            # Configuration du contexte
            sns.set_context("paper", rc={
                'font.size': 12,
                'axes.labelsize': 14,
                'axes.titlesize': 16,
                'xtick.labelsize': 11,
                'ytick.labelsize': 11,
                'legend.fontsize': 11,
                'legend.title_fontsize': 12
            }, font_scale=1.2)
            
            # Configuration matplotlib
            plt.rcParams.update({
                # Police et taille
                'font.family': 'serif',
                'font.serif': ['Computer Modern Roman', 'DejaVu Serif', 'Times New Roman'],
                'font.sans-serif': ['Helvetica', 'Arial'],
                
                # Épaisseur des lignes
                'lines.linewidth': 2,
                'lines.markersize': 8,
                'lines.markeredgewidth': 1.5,
                
                # Style des axes
                'axes.spines.top': False,
                'axes.spines.right': False,
                'axes.labelweight': 'bold',
                'axes.titleweight': 'bold',
                'axes.titlepad': 20,
                
                # Grille
                'grid.linewidth': 0.5,
                
                # Légende
                'legend.frameon': True,
                'legend.framealpha': 0.8,
                'legend.edgecolor': colors['grid'],
                'legend.fancybox': True,
                
                # Figure
                'figure.figsize': (12, 8),
                'figure.dpi': 300,
                'figure.facecolor': 'white',
                'figure.titlesize': 18,
                'figure.titleweight': 'bold',
                
                # Marges
                'figure.subplot.top': 0.85,
                'figure.subplot.right': 0.95,
                'figure.subplot.left': 0.1,
                'figure.subplot.bottom': 0.15,
                
                # Sauvegarde
                'savefig.dpi': 300,
                'savefig.bbox': 'tight',
                'savefig.pad_inches': 0.2,
                
                # Animations
                'animation.embed_limit': 100,
                
                # Style général
                'axes.prop_cycle': plt.cycler('color', [
                    colors['primary'],
                    colors['secondary'],
                    colors['accent'],
                    colors['success'],
                    colors['warning']
                ])
            })
            
            # Retourner les couleurs pour utilisation dans les graphiques
            return colors
            
        except ImportError:
            # Fallback si seaborn n'est pas disponible
            plt.style.use('default')
            return {
                'primary': 'blue',
                'secondary': 'red',
                'accent': 'green',
                'neutral': 'gray',
                'success': 'darkgreen',
                'warning': 'orange',
                'background': 'whitesmoke',
                'grid': 'lightgray'
            }

    def create_visualizations(self, data: Dict[str, Any], filename_prefix: str) -> Dict[str, Path]:
        """Crée toutes les visualisations nécessaires pour le rapport."""
        # Vérifier le cache
        cache_key = self.generate_cache_key(data)
        cached_images = self.get_cached_images(cache_key)
        
        # Si toutes les images sont en cache, les réutiliser
        if len(cached_images) == 4:  # Nous avons 4 types d'images
            print("Utilisation des images en cache...")
            images = {}
            for img_type, cached_path in cached_images.items():
                new_path = self.output_dir / f"{filename_prefix}_{img_type}.png"
                if not new_path.exists():  # Éviter la copie si le fichier existe déjà
                    import shutil
                    shutil.copy2(cached_path, new_path)
                images[img_type] = new_path
            return images
            
        print("Génération de nouvelles images...")
        images = {}
        
        # Configuration du style
        colors = self.setup_plotting_style()
        
        # 1. Distribution des Différences et Probabilités
        plt.figure(figsize=(12, 8))
        
        # Préparation des données
        differences = list(map(int, data['differences'].keys()))
        frequencies = list(data['differences'].values())
        total_samples = sum(frequencies)
        probabilities = [f/total_samples for f in frequencies]
        
        # Calcul des statistiques
        mean_freq = np.mean(frequencies)
        std_freq = np.std(frequencies)
        max_freq = max(frequencies)
        max_freq_idx = frequencies.index(max_freq)
        
        # Création du graphique principal
        ax = plt.gca()
        bars = plt.bar(range(len(differences)), frequencies, 
                      color=colors['accent'], edgecolor=colors['primary'],
                      alpha=0.7, label='Fréquence observée')
        
        # Ligne de tendance et intervalle de confiance
        x = np.arange(len(differences))
        z = np.polyfit(x, frequencies, 2)
        p = np.poly1d(z)
        plt.plot(x, p(x), "--", color=colors['secondary'], alpha=0.8,
                label='Tendance polynomiale', linewidth=2)
        
        # Intervalle de confiance (±2σ)
        plt.fill_between(x,
                       [mean_freq - 2*std_freq]*len(differences),
                       [mean_freq + 2*std_freq]*len(differences),
                       color=colors['neutral'], alpha=0.2,
                       label='Intervalle de confiance (±2σ)')
        
        # Ligne de référence (distribution uniforme)
        plt.axhline(y=mean_freq, color=colors['success'],
                   linestyle='--', alpha=0.5,
                   label='Moyenne théorique (uniforme)')
        
        # Annotations
        plt.annotate(f'Pic maximal\n{max_freq} occurrences\n({(max_freq/total_samples):.1%})',
                    xy=(max_freq_idx, max_freq),
                    xytext=(10, 10), textcoords='offset points',
                    ha='center', va='bottom',
                    bbox=dict(boxstyle='round,pad=0.5',
                            fc=colors['warning'], alpha=0.5,
                            ec=colors['primary']),
                    arrowprops=dict(arrowstyle='->',
                                  connectionstyle='arc3,rad=0.2',
                                  color=colors['primary']))
        
        # Statistiques dans un encadré
        stats_text = (f'Statistiques:\n'
                     f'Nombre total: {total_samples}\n'
                     f'Moyenne: {mean_freq:.1f}\n'
                     f'Écart-type: {std_freq:.1f}\n'
                     f'Entropie: {self.calculate_entropy(data["differences"]):.2f} bits')
        plt.text(0.95, 0.95, stats_text,
                transform=ax.transAxes,
                verticalalignment='top',
                horizontalalignment='right',
                bbox=dict(boxstyle='round',
                        facecolor=colors['background'],
                        edgecolor=colors['primary'],
                        alpha=0.8))

        # Personnalisation du graphique
        plt.title('Distribution des Différences de Sortie\n'
                 'Analyse de la propagation des différences', 
                 pad=20, color=colors['primary'])
        plt.xlabel('Valeur de la Différence (en hexadécimal)',
                  color=colors['primary'])
        plt.ylabel('Nombre d\'Occurrences',
                  color=colors['primary'])
        
        # Amélioration des étiquettes de l'axe x
        plt.xticks(range(len(differences)), 
                  [f'0x{d:04x}' for d in differences], 
                  rotation=45, ha='right')
        
        # Grille et légende
        plt.grid(True, linestyle='--', alpha=0.3)
        plt.legend(loc='upper left', bbox_to_anchor=(1.02, 1),
                  borderaxespad=0)
        
        # Ajustement de la mise en page
        plt.tight_layout()
        plt.savefig(self.output_dir / f"{filename_prefix}_probability.png",
                   dpi=300, bbox_inches='tight')
        images['probability'] = self.output_dir / f"{filename_prefix}_probability.png"
        plt.close()

        # 2. Analyse des Probabilités
        plt.figure(figsize=(12, 8))
        
        # Préparation des données
        prob = data['probability']
        random_prob = 1.0 / (1 << data['block_size'])
        bias_factor = prob / random_prob
        probs = [random_prob, prob]
        labels = ['Probabilité\nAléatoire\n(Théorique)', 'Probabilité\nObservée\n(Attaque)']
        
        # Création des barres
        ax = plt.gca()
        bars = plt.bar(labels, probs, 
                      color=[colors['neutral'], colors['accent']],
                      edgecolor=[colors['primary'], colors['primary']],
                      alpha=0.7)
        
        # Ajout des valeurs sur les barres
        for bar in bars:
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width()/2., height,
                    f'{height:.2e}',
                    ha='center', va='bottom',
                    bbox=dict(boxstyle='round,pad=0.5', 
                             fc=colors['background'],
                             ec=colors['primary'],
                             alpha=0.8))
        
        # Ligne de référence et zone de significativité
        plt.axhline(y=random_prob, color=colors['secondary'],
                   linestyle='--', alpha=0.5,
                   label='Seuil de l\'aléatoire')
        
        # Zone de significativité statistique
        significant_region = np.array([random_prob * (1 + x/100) for x in range(-10, 11)])
        x = np.array([-0.5, 1.5])
        plt.fill_between(x, [significant_region.min()]*2, [significant_region.max()]*2,
                        color=colors['neutral'], alpha=0.1,
                        label='Zone de non-significativité (±10%)')
        
        # Calcul et affichage du facteur de biais
        bias_text = (f'Facteur de biais : {bias_factor:.1f}x\n'
                    f'L\'attaque est {bias_factor:.1f} fois\n'
                    f'plus efficace que le hasard\n\n'
                    f'p-value : {self.calculate_p_value(data["differences"]):.2e}')
        plt.text(0.5, prob/2, bias_text,
                ha='center', va='center',
                bbox=dict(boxstyle='round,pad=0.5', 
                         fc=colors['warning'],
                         ec=colors['primary'],
                         alpha=0.5))

        # Personnalisation du graphique
        plt.title('Comparaison des Probabilités\n'
                 'Efficacité de l\'attaque vs Distribution Aléatoire', 
                 pad=20, color=colors['primary'])
        plt.ylabel('Probabilité (échelle logarithmique)',
                  color=colors['primary'])
        plt.yscale('log')
        
        # Grille logarithmique
        plt.grid(True, which="both", ls="-", alpha=0.2)
        plt.grid(True, which="minor", ls=":", alpha=0.1)
        
        # Légende
        plt.legend(loc='upper left', bbox_to_anchor=(1.02, 1),
                  borderaxespad=0)
        
        # Ajustement de la mise en page
        plt.tight_layout()
        plt.savefig(self.output_dir / f"{filename_prefix}_probability.png",
                   dpi=300, bbox_inches='tight')
        images['probability'] = self.output_dir / f"{filename_prefix}_probability.png"
        plt.close()

        # 3. Heatmap de Corrélation
        plt.figure(figsize=(12, 10))
        
        # Calcul de la matrice de corrélation
        corr_matrix = self.calculate_correlation_matrix(data)
        
        # Création de la heatmap avec une palette personnalisée
        ax = plt.gca()
        im = plt.imshow(corr_matrix, 
                       cmap='RdYlBu_r',  # Rouge-Jaune-Bleu inversé
                       aspect='auto')
        
        # Ajout d'une barre de couleur personnalisée
        cbar = plt.colorbar(im, label='Coefficient de corrélation')
        cbar.ax.tick_params(labelsize=10)
        
        # Étiquettes
        differences = list(map(int, data['differences'].keys()))
        plt.xticks(range(len(differences)), 
                  [f'0x{d:04x}' for d in differences],
                  rotation=45, ha='right')
        plt.yticks(range(len(differences)), 
                  [f'0x{d:04x}' for d in differences])
        
        # Titre et labels avec style personnalisé
        plt.title('Matrice de Corrélation des Différences\n'
                 'Analyse des relations entre les différences',
                 pad=20, color=colors['primary'])
        plt.xlabel('Différence', color=colors['primary'])
        plt.ylabel('Différence', color=colors['primary'])
        
        # Ajout d'un cadre autour de la heatmap
        for spine in ax.spines.values():
            spine.set_visible(True)
            spine.set_color(colors['primary'])
            spine.set_linewidth(1.5)

        # Ajustement de la mise en page
        plt.tight_layout()
        plt.savefig(self.output_dir / f"{filename_prefix}_correlation.png",
                   dpi=300, bbox_inches='tight')
        images['correlation'] = self.output_dir / f"{filename_prefix}_correlation.png"
        plt.close()

        # 4. Boxplot des Distributions
        plt.figure(figsize=(12, 8))
        
        # Préparation des données
        normalized_diffs, frequencies = self.prepare_boxplot_data(data)
        
        # Création du boxplot avec couleurs personnalisées
        bp = plt.boxplot([normalized_diffs, frequencies], 
                        labels=['Différences\nNormalisées', 'Fréquences\nRelatives'],
                        patch_artist=True,
                        medianprops=dict(color=colors['primary'], linewidth=1.5),
                        flierprops=dict(marker='o', markerfacecolor=colors['warning'],
                                      markeredgecolor=colors['primary'],
                                      markersize=8))
        
        # Personnalisation des couleurs des boîtes
        box_colors = [colors['accent'], colors['secondary']]
        for patch, color in zip(bp['boxes'], box_colors):
            patch.set_facecolor(color)
            patch.set_alpha(0.7)
            patch.set_edgecolor(colors['primary'])
        
        # Ajout des points individuels
        for i, (x, y) in enumerate([(1, normalized_diffs), (2, frequencies)]):
            plt.scatter([x] * len(y), y, 
                       alpha=0.5, color=colors['primary'], 
                       label='Observations' if i == 0 else "")
        
        # Statistiques dans un encadré avec style personnalisé
        stats_text = (f'Statistiques:\n'
                     f'Différences :\n'
                     f'  Médiane: {np.median(normalized_diffs):.3f}\n'
                     f'  IQR: {np.percentile(normalized_diffs, 75) - np.percentile(normalized_diffs, 25):.3f}\n'
                     f'Fréquences :\n'
                     f'  Médiane: {np.median(frequencies):.3f}\n'
                     f'  IQR: {np.percentile(frequencies, 75) - np.percentile(frequencies, 25):.3f}')
        
        plt.text(0.95, 0.95, stats_text,
                transform=plt.gca().transAxes,
                verticalalignment='top',
                horizontalalignment='right',
                bbox=dict(boxstyle='round',
                         facecolor=colors['background'],
                         edgecolor=colors['primary'],
                         alpha=0.8))

        # Personnalisation du graphique
        plt.title('Distribution des Différences et Fréquences\n'
                 'Analyse de la dispersion des valeurs',
                 pad=20, color=colors['primary'])
        plt.ylabel('Valeur normalisée', color=colors['primary'])
        plt.grid(True, linestyle='--', alpha=0.3)

        # Ajustement de la mise en page
        plt.tight_layout()
        plt.savefig(self.output_dir / f"{filename_prefix}_boxplot.png",
                   dpi=300, bbox_inches='tight')
        images['boxplot'] = self.output_dir / f"{filename_prefix}_boxplot.png"
        plt.close()

        # 5. Graphique de Convergence
        plt.figure(figsize=(12, 8))
        
        # Calcul des données de convergence
        samples, probabilities = self.calculate_convergence_data(data)
        
        # Création du graphique avec style personnalisé
        plt.plot(samples, probabilities, '-',
                color=colors['accent'],
                linewidth=2.5, alpha=0.7,
                label='Probabilité observée')
        
        # Ligne de convergence théorique
        theoretical_prob = data['probability']
        plt.axhline(y=theoretical_prob, 
                   color=colors['secondary'],
                   linestyle='--', alpha=0.5,
                   label='Probabilité théorique')
        
        # Zone de convergence avec couleur personnalisée
        plt.fill_between(samples,
                        [theoretical_prob * 0.9] * len(samples),
                        [theoretical_prob * 1.1] * len(samples),
                        color=colors['neutral'], alpha=0.2,
                        label='Zone de convergence (±10%)')
        
        # Annotations avec style personnalisé
        final_prob = probabilities[-1]
        plt.annotate(f'Probabilité finale: {final_prob:.2e}',
                    xy=(samples[-1], final_prob),
                    xytext=(-100, 20), textcoords='offset points',
                    ha='right', va='bottom',
                    bbox=dict(boxstyle='round,pad=0.5',
                            fc=colors['warning'],
                            ec=colors['primary'],
                            alpha=0.5),
                    arrowprops=dict(arrowstyle='->',
                                  connectionstyle='arc3,rad=0.2',
                                  color=colors['primary']))
        
        # Personnalisation des axes logarithmiques
        plt.xscale('log')
        plt.yscale('log')
        
        # Grille logarithmique personnalisée
        plt.grid(True, which="both", ls="-", alpha=0.2,
                color=colors['grid'])
        plt.grid(True, which="minor", ls=":", alpha=0.1,
                color=colors['grid'])
        
        # Titre et labels avec style personnalisé
        plt.title('Convergence de l\'Attaque\n'
                 'Évolution de la probabilité avec le nombre d\'échantillons',
                 pad=20, color=colors['primary'])
        plt.xlabel('Nombre d\'échantillons', color=colors['primary'])
        plt.ylabel('Probabilité', color=colors['primary'])
        
        # Légende avec style personnalisé
        plt.legend(loc='upper left',
                  bbox_to_anchor=(1.02, 1),
                  borderaxespad=0,
                  frameon=True,
                  facecolor=colors['background'],
                  edgecolor=colors['primary'])
        
        # Ajustement de la mise en page
        plt.tight_layout()
        plt.savefig(self.output_dir / f"{filename_prefix}_convergence.png",
                   dpi=300, bbox_inches='tight')
        images['convergence'] = self.output_dir / f"{filename_prefix}_convergence.png"
        plt.close()
        
        # Sauvegarder dans le cache
        self.save_to_cache(cache_key, images)
        
        return images

    def get_feistel_introduction(self) -> str:
        """Obtient une introduction sur les schémas de Feistel via OpenAI."""
        prompt = """Rédigez une introduction claire et concise sur les schémas de Feistel en cryptographie, incluant :
1. Leur principe de fonctionnement
2. Leurs caractéristiques principales
3. Leur importance dans la cryptographie moderne
4. Quelques exemples d'algorithmes connus utilisant ce schéma"""

        try:
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "Vous êtes un expert en cryptographie spécialisé dans les schémas de Feistel."},
                    {"role": "user", "content": prompt}
                ]
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Erreur lors de la génération de l'introduction: {str(e)}"

    def analyze_attack_results(self, data: Dict[str, Any]) -> str:
        """Analyse les résultats de l'attaque différentielle."""
        bias_factor = data['probability'] / (1.0 / (1 << data['block_size']))
        prompt = f"""Analysez en détail les résultats suivants d'une attaque différentielle sur un schéma de Feistel :

Configuration :
- Nombre de tours : {data['num_rounds']}
- Taille de bloc : {data['block_size']} bits

Résultats :
- Différence en entrée : {data['input_diff']:08x}
- Différence en sortie : {data['output_diff']:08x}
- Probabilité observée : {data['probability']:.2%}
- Probabilité aléatoire théorique : {1.0 / (1 << data['block_size']):.2e}

Fournissez :
1. Une explication détaillée de l'attaque différentielle réalisée
2. Une analyse des biais observés et leur signification
3. Une évaluation de la vulnérabilité du chiffrement
4. Des recommandations concrètes pour renforcer la sécurité"""

        try:
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "Vous êtes un expert en cryptanalyse spécialisé dans l'analyse des attaques différentielles."},
                    {"role": "user", "content": prompt}
                ]
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Erreur lors de l'analyse des résultats: {str(e)}"

    def get_graph_explanations(self, data: Dict[str, Any]) -> str:
        """Génère des explications détaillées des graphiques pour les non-experts."""
        bias_factor = data['probability'] / (1.0 / (1 << data['block_size']))
        
        prompt = f"""Expliquez de manière simple et accessible les deux graphiques suivants d'une analyse cryptographique, en évitant le jargon technique autant que possible :

1. Le premier graphique (Distribution des Différences) :
   - Montre comment les différences se répartissent
   - Le pic maximal indique la différence la plus fréquente
   - Plus les pics sont marqués, plus l'attaque est efficace

2. Le second graphique (Comparaison des Probabilités) :
   - Compare la probabilité observée lors de l'attaque avec celle du hasard
   - Montre que l'attaque est {bias_factor:.1f} fois plus efficace que le hasard
   - Utilise une échelle logarithmique pour mieux visualiser les différences

Expliquez :
1. Ce que ces graphiques signifient pour un public non technique
2. Pourquoi ces résultats sont importants
3. Ce qu'ils nous apprennent sur la sécurité du système
4. Utilisez des analogies simples pour faciliter la compréhension"""

        try:
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "Vous êtes un expert en vulgarisation scientifique, spécialisé dans l'explication de concepts cryptographiques complexes à un public non technique."},
                    {"role": "user", "content": prompt}
                ]
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Erreur lors de la génération des explications: {str(e)}"

    def calculate_entropy(self, differences: Dict[int, int]) -> float:
        """Calcule l'entropie de Shannon de la distribution des différences."""
        total = sum(differences.values())
        probabilities = [count/total for count in differences.values()]
        return -sum(p * np.log2(p) for p in probabilities if p > 0)

    def calculate_uniformity(self, differences: Dict[int, int]) -> float:
        """Calcule l'uniformité relative de la distribution."""
        total = sum(differences.values())
        expected = total / len(differences)
        max_deviation = max(abs(count - expected) for count in differences.values())
        return 1 - (max_deviation / expected)

    def chi_square_test(self, differences: Dict[int, int]) -> float:
        """Calcule la statistique du test du χ²."""
        total = sum(differences.values())
        expected = total / len(differences)
        chi_square = sum((count - expected)**2 / expected for count in differences.values())
        return chi_square

    def calculate_p_value(self, differences: Dict[int, int]) -> float:
        """Calcule la p-value du test du χ²."""
        chi_square = self.chi_square_test(differences)
        df = len(differences) - 1
        return 1 - chi2.cdf(chi_square, df)

    def calculate_correlation_matrix(self, data: Dict[str, Any]) -> np.ndarray:
        """Calcule la matrice de corrélation entre les différences."""
        differences = list(map(int, data['differences'].keys()))
        n = len(differences)
        corr_matrix = np.zeros((n, n))
        
        # Convertir les différences en format binaire
        bin_diffs = [bin(d)[2:].zfill(data["block_size"]) for d in differences]
        
        # Calculer les corrélations bit à bit
        for i in range(n):
            for j in range(n):
                # Compter les bits différents
                xor = bin(int(bin_diffs[i], 2) ^ int(bin_diffs[j], 2))[2:]
                hamming_dist = xor.count('1')
                # Normaliser la distance de Hamming
                corr_matrix[i, j] = 1 - (hamming_dist / data["block_size"])
        
        return corr_matrix

    def prepare_boxplot_data(self, data: Dict[str, Any]) -> Tuple[List[float], List[float]]:
        """Prépare les données pour le boxplot."""
        # Convertir les différences en valeurs normalisées
        differences = list(map(int, data['differences'].keys()))
        max_diff = max(differences)
        normalized_diffs = [d/max_diff for d in differences]
        
        # Calculer les fréquences relatives
        total = sum(data['differences'].values())
        frequencies = [count/total for count in data['differences'].values()]
        
        return normalized_diffs, frequencies

    def calculate_convergence_data(self, data: Dict[str, Any]) -> Tuple[List[int], List[float]]:
        """Calcule les données de convergence de l'attaque."""
        total_samples = sum(data['differences'].values())
        samples = list(range(100, total_samples + 1, total_samples // 100))
        
        # Simuler la convergence en accumulant les échantillons
        probabilities = []
        max_diff = max(data['differences'], key=data['differences'].get)
        
        for n in samples:
            # Simuler n échantillons
            scaled_counts = {k: int(v * n/total_samples) 
                           for k, v in data['differences'].items()}
            prob = scaled_counts[max_diff] / n if n > 0 else 0
            probabilities.append(prob)
        
        return samples, probabilities

    def generate_report(self, data: Dict[str, Any]) -> Path:
        """Génère un rapport complet avec visualisations et analyses."""
        # Créer le répertoire de sortie s'il n'existe pas
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Générer un préfixe unique pour les fichiers
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename_prefix = f"report_{timestamp}"
        
        # Créer les visualisations
        images = self.create_visualizations(data, filename_prefix)
        
        # Générer le rapport Word
        report_path = self.generate_word_report(data, filename_prefix, images)
        
        return report_path

    def clean_html(self, html_text: str) -> str:
        """Nettoie les balises HTML du texte."""
        clean = re.compile('<.*?>')
        return re.sub(clean, '', html_text)

    def generate_word_report(self, data: Dict[str, Any], filename_prefix: str, images: Dict[str, Path]) -> Path:
        """Génère un rapport Word formaté avec les résultats de l'analyse."""
        doc = Document()
        
        # Styles personnalisés
        styles = doc.styles
        
        # Style pour les titres
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.size = Pt(24)
        title_font.bold = True
        title_font.color.rgb = RGBColor(44, 62, 80)  # Bleu foncé
        
        # Style pour les sous-titres
        heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_font = heading_style.font
        heading_font.size = Pt(16)
        heading_font.bold = True
        heading_font.color.rgb = RGBColor(52, 152, 219)  # Bleu clair
        
        # Titre principal
        title = doc.add_paragraph("Rapport d'Analyse Cryptographique", style='CustomTitle')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Sous-titre avec date
        subtitle = doc.add_paragraph(f"Analyse Différentielle d'un Schéma de Feistel à {data['num_rounds']} Tours", style='CustomHeading')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Date: {datetime.now().strftime('%d/%m/%Y')}")
        
        # Table des matières
        doc.add_paragraph("Table des Matières", style='CustomHeading')
        toc_entries = [
            "1. Résumé Exécutif",
            "2. Introduction",
            "3. Méthodologie",
            "4. Résultats de l'Analyse",
            "5. Visualisations",
            "6. Conclusions",
            "7. Recommandations"
        ]
        for entry in toc_entries:
            doc.add_paragraph(entry, style='List Number')
        doc.add_page_break()
        
        # Résumé exécutif
        doc.add_paragraph("1. Résumé Exécutif", style='CustomHeading')
        exec_summary = self.generate_executive_summary(data)
        doc.add_paragraph(self.clean_html(markdown.markdown(exec_summary)))
        
        # Introduction
        doc.add_paragraph("2. Introduction", style='CustomHeading')
        doc.add_paragraph("""Cette analyse porte sur l'évaluation de la sécurité d'un schéma de Feistel 
        face aux attaques différentielles. L'objectif est d'identifier d'éventuelles faiblesses dans 
        la structure cryptographique et de proposer des recommandations d'amélioration.""")
        
        # Méthodologie
        total_pairs = sum(data['differences'].values())
        doc.add_paragraph("3. Méthodologie", style='CustomHeading')
        doc.add_paragraph(f"""L'analyse a été réalisée sur un schéma de Feistel avec les paramètres suivants :
        • Taille de bloc : {data['block_size']} bits
        • Nombre de tours : {data['num_rounds']}
        • Nombre de paires analysées : {total_pairs:,}""")
        
        # Résultats
        doc.add_paragraph("4. Résultats de l'Analyse", style='CustomHeading')
        results_text = self.generate_analysis_text(data)
        doc.add_paragraph(self.clean_html(markdown.markdown(results_text)))
        
        # Visualisations
        doc.add_paragraph("5. Visualisations", style='CustomHeading')
        # Ajouter les images avec légendes
        image_captions = {
            'probability': "Distribution des Différences et Probabilités",
            'correlation': "Matrice de Corrélation des Différences",
            'boxplot': "Analyse de la Distribution par Boxplot",
            'convergence': "Convergence de l'Attaque"
        }
        
        for img_type, img_path in images.items():
            if img_path.exists():
                doc.add_picture(str(img_path), width=Inches(6))
                caption = image_captions.get(img_type, img_type.title())
                caption_par = doc.add_paragraph(caption)
                caption_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()  # Espace après l'image
        
        # Conclusions
        doc.add_paragraph("6. Conclusions", style='CustomHeading')
        conclusions = self.generate_conclusions(data)
        doc.add_paragraph(self.clean_html(markdown.markdown(conclusions)))
        
        # Recommandations
        doc.add_paragraph("7. Recommandations", style='CustomHeading')
        recommendations = self.generate_recommendations(data)
        doc.add_paragraph(self.clean_html(markdown.markdown(recommendations)))
        
        # Sauvegarder le document
        output_path = self.output_dir / f"{filename_prefix}_report.docx"
        doc.save(str(output_path))
        return output_path

    def generate_executive_summary(self, data: Dict[str, Any]) -> str:
        """Génère un résumé exécutif pour le rapport."""
        return f"""Cette analyse a révélé des biais statistiques significatifs dans la distribution des différences de sortie du schéma de Feistel. 
        La probabilité observée est de {data['probability']:.2%}, ce qui est {data['probability'] / (1.0 / (1 << data['block_size'])):.1f} fois plus élevée que la probabilité aléatoire théorique. 
        Ces résultats suggèrent que le schéma de Feistel est vulnérable aux attaques différentielles et que des mesures de sécurité supplémentaires sont nécessaires."""

    def generate_analysis_text(self, data: Dict[str, Any]) -> str:
        """Génère le texte d'analyse pour le rapport."""
        total_pairs = sum(data['differences'].values())
        return f"""L'analyse a été réalisée sur un schéma de Feistel avec les paramètres suivants :
        • Taille de bloc : {data['block_size']} bits
        • Nombre de tours : {data['num_rounds']}
        • Nombre de paires analysées : {total_pairs:,}
        
        Les résultats de l'analyse montrent que la distribution des différences de sortie est loin d'être uniforme. 
        La probabilité observée est de {data['probability']:.2%}, ce qui est {data['probability'] / (1.0 / (1 << data['block_size'])):.1f} fois plus élevée que la probabilité aléatoire théorique. 
        Ces résultats suggèrent que le schéma de Feistel est vulnérable aux attaques différentielles."""

    def generate_conclusions(self, data: Dict[str, Any]) -> str:
        """Génère les conclusions pour le rapport."""
        return f"""En conclusion, cette analyse a révélé des biais statistiques significatifs dans la distribution des différences de sortie du schéma de Feistel. 
        Ces résultats suggèrent que le schéma de Feistel est vulnérable aux attaques différentielles et que des mesures de sécurité supplémentaires sont nécessaires. 
        Il est recommandé de prendre des mesures pour renforcer la sécurité du schéma de Feistel, telles que l'augmentation du nombre de tours ou la modification de la fonction de tour."""

    def generate_recommendations(self, data: Dict[str, Any]) -> str:
        """Génère les recommandations pour le rapport."""
        return f"""Il est recommandé de prendre les mesures suivantes pour renforcer la sécurité du schéma de Feistel :
        • Augmenter le nombre de tours
        • Modifier la fonction de tour
        • Ajouter des mesures de sécurité supplémentaires, telles que des mécanismes de détection d'intrusion ou des systèmes de gestion des clés."""
